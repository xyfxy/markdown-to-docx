import pytest
from pathlib import Path
import sys
import os
import docx # For reading DOCX files
from click.testing import CliRunner

# Adjust sys.path to ensure main and its dependencies can be imported
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from main import markdown_to_docx_cli # Assuming the command in main.py is markdown_to_docx_cli

# --- Fixtures ---

@pytest.fixture
def create_md_file(tmp_path: Path):
    """
    Pytest fixture to create a temporary Markdown file with given content.
    `tmp_path` is a built-in pytest fixture providing a temporary directory.
    """
    files_created = []
    def _create_file(content: str, filename: str = "test_input.md"):
        file_path = tmp_path / filename
        file_path.write_text(content, encoding='utf-8')
        files_created.append(file_path)
        return str(file_path)
    yield _create_file
    # Optional: cleanup files if needed, though tmp_path usually handles it.
    # for f_path in files_created:
    #     try:
    #         os.remove(f_path)
    #     except OSError:
    #         pass

@pytest.fixture
def runner():
    """Provides a CliRunner instance."""
    return CliRunner()

# --- Helper for DOCX text extraction ---
def get_paragraph_texts(docx_path: str) -> list[str]:
    """Extracts non-empty text from all paragraphs in a DOCX file."""
    try:
        doc = docx.Document(docx_path)
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    except Exception as e:
        pytest.fail(f"Failed to read DOCX file {docx_path}: {e}")
        return []

# --- Integration Tests ---

def test_simple_markdown_conversion(create_md_file, tmp_path, runner):
    """Test basic conversion of headings, paragraphs, and simple lists."""
    md_content = """
# Main Title

This is a paragraph.

- Item 1
- Item 2

1. Ordered A
2. Ordered B
"""
    md_file = create_md_file(md_content)
    output_docx = str(tmp_path / "simple_output.docx")

    result = runner.invoke(markdown_to_docx_cli, [md_file, output_docx])

    assert result.exit_code == 0, f"CLI Error: {result.output}"
    assert "Successfully converted" in result.output
    assert os.path.exists(output_docx)

    texts = get_paragraph_texts(output_docx)
    
    # Note: DocxWriter adds prefixes to headings (H2-H5) and empty lines after H1.
    # Main Title (H1) does not have a prefix from DocxWriter's _add_main_heading.
    assert "Main Title" in texts[0] # H1
    # _add_main_heading adds 2 empty paragraphs, so these won't be in `texts` due to strip()
    
    assert "This is a paragraph." in texts # Order might vary due to empty lines
    
    # List items - DocxWriter adds prefixes like "• " and "1. "
    # The exact paragraph index might be tricky due to potential empty paragraphs from Markdown
    # or spacing paragraphs from DocxWriter. Using 'in texts' is more robust.
    assert any("• Item 1" in t for t in texts)
    assert any("• Item 2" in t for t in texts)
    assert any("1. Ordered A" in t for t in texts)
    assert any("2. Ordered B" in t for t in texts)

def test_conversion_with_cli_options(create_md_file, tmp_path, runner):
    """Test conversion with --font-size and --line-spacing CLI options."""
    md_content = "A simple paragraph for style testing."
    md_file = create_md_file(md_content)
    output_docx = str(tmp_path / "styled_output.docx")
    
    custom_font_size = "12.0" # Corresponds to 小四
    custom_line_spacing = "20.5"

    result = runner.invoke(markdown_to_docx_cli, [
        md_file, output_docx,
        '--font-size', custom_font_size,
        '--line-spacing', custom_line_spacing
    ])

    assert result.exit_code == 0, f"CLI Error: {result.output}"
    assert "Successfully converted" in result.output
    assert f"Custom body font size: {custom_font_size} pt" in result.output
    assert f"Custom line spacing: {custom_line_spacing} pt" in result.output
    assert os.path.exists(output_docx)

    # Direct DOCX style validation is complex for integration tests.
    # We primarily test that the CLI options are accepted and the conversion succeeds.
    # Unit tests for DocxWriter should confirm these values are used.
    # Optionally, a very basic check if python-docx allows easy access:
    try:
        doc = docx.Document(output_docx)
        # Check Normal style if it reflects the changes (it should)
        # This is a bit of a white-box test for an integration test.
        normal_style = doc.styles['Normal']
        assert normal_style.font.size == Pt(float(custom_font_size))
        assert normal_style.paragraph_format.line_spacing == Pt(float(custom_line_spacing))
        
        # Check the actual paragraph if it inherited correctly
        # (DocxWriter applies styles run by run for body text)
        para_texts = get_paragraph_texts(output_docx)
        assert md_content in para_texts

        # Find the paragraph and check a run's font size
        # This assumes the first non-empty paragraph is the one.
        first_content_p = None
        for p in doc.paragraphs:
            if p.text.strip() == md_content:
                first_content_p = p
                break
        
        assert first_content_p is not None, "Content paragraph not found"
        if first_content_p.runs:
            # Check the font size of the first run in the paragraph
            assert first_content_p.runs[0].font.size == Pt(float(custom_font_size))

    except Exception as e:
        pytest.fail(f"Failed to perform basic style check on DOCX {output_docx}: {e}")


def test_conversion_with_utf8_content(create_md_file, tmp_path, runner):
    """Test conversion of Markdown with UTF-8 (Chinese) characters."""
    md_content = """
# 中文标题

这是一段中文内容。

- 列表项一
- 列表项二
"""
    md_file = create_md_file(md_content, "utf8_input.md")
    output_docx = str(tmp_path / "utf8_output.docx")

    result = runner.invoke(markdown_to_docx_cli, [md_file, output_docx])

    assert result.exit_code == 0, f"CLI Error: {result.output}"
    assert os.path.exists(output_docx)

    texts = get_paragraph_texts(output_docx)
    assert "中文标题" in texts[0] # H1
    assert "这是一段中文内容。" in texts
    assert any("• 列表项一" in t for t in texts) # DocxWriter adds "• " prefix
    assert any("• 列表项二" in t for t in texts)


def test_list_conversion_text_and_prefix(create_md_file, tmp_path, runner):
    """Test detailed list conversion including prefixes and nesting."""
    md_content = """
- Level 0 Unordered
  1. Level 1 Ordered A
  2. Level 1 Ordered B
     - Level 2 Unordered B.1
- Level 0 Unordered Again
"""
    md_file = create_md_file(md_content)
    output_docx = str(tmp_path / "list_output.docx")

    result = runner.invoke(markdown_to_docx_cli, [md_file, output_docx])
    assert result.exit_code == 0, f"CLI Error: {result.output}"
    texts = get_paragraph_texts(output_docx)

    # Based on UNORDERED_LIST_BULLETS = ["• ", "◦ ", "▪ "] and ORDERED_LIST_PREFIX_FORMAT = "{}. "
    # And how main.py calculates level for prefix.
    assert any("• Level 0 Unordered" in t for t in texts)
    assert any("1. Level 1 Ordered A" in t for t in texts) # Prefixed by main.py list logic
    assert any("2. Level 1 Ordered B" in t for t in texts)
    assert any("◦ Level 2 Unordered B.1" in t for t in texts) # Nested unordered, should use "◦ "
    assert any("• Level 0 Unordered Again" in t for t in texts)


def test_error_handling_input_file_not_found(tmp_path, runner):
    """Test CLI error handling for a non-existent input Markdown file."""
    non_existent_md = str(tmp_path / "non_existent.md")
    output_docx = str(tmp_path / "error_output.docx")

    result = runner.invoke(markdown_to_docx_cli, [non_existent_md, output_docx])

    assert result.exit_code != 0 # Expect non-zero exit code for error
    # Click itself handles "does not exist" for type=click.Path(exists=True)
    assert f"Error: Invalid value for 'INPUT_MD': Path '{non_existent_md}' does not exist." in result.output or \
           f"Error: File not found - {non_existent_md}" in result.output # If main.py's own check triggers

def test_recipient_heuristic(create_md_file, tmp_path, runner):
    """Test the heuristic for identifying a recipient line."""
    md_content = """
# Document Title

尊敬的张三先生：

这是文件的主要内容。
"""
    # The heuristic in main.py checks if the FIRST paragraph (after any H1) ends with a colon.
    # DocxWriter's _add_main_heading adds two empty lines after H1.
    # Then the recipient line. Then the main content.
    
    md_file = create_md_file(md_content)
    output_docx = str(tmp_path / "recipient_output.docx")

    result = runner.invoke(markdown_to_docx_cli, [md_file, output_docx])
    assert result.exit_code == 0, f"CLI Error: {result.output}"
    assert "Processed first paragraph as recipient" in result.output # Check log output

    texts = get_paragraph_texts(output_docx)
    
    # Expected: Title, then Recipient (with colon added by DocxWriter), then Content.
    # The empty lines from _add_main_heading are stripped by get_paragraph_texts.
    assert "Document Title" == texts[0]
    assert "尊敬的张三先生：" == texts[1] # DocxWriter.add_recipient appends "："
    assert "这是文件的主要内容。" == texts[2]

def test_empty_markdown_file(create_md_file, tmp_path, runner):
    """Test conversion of an empty markdown file."""
    md_file = create_md_file("") # Empty content
    output_docx = str(tmp_path / "empty_output.docx")

    result = runner.invoke(markdown_to_docx_cli, [md_file, output_docx])
    assert result.exit_code == 0, f"CLI Error: {result.output}"
    assert "Markdown parsing resulted in no elements" in result.output # Check log/print
    assert os.path.exists(output_docx)

    doc = docx.Document(output_docx)
    # An empty markdown might result in a docx with just default empty paragraph, or nothing if parser is robust.
    # The current logic saves an empty doc if no elements are parsed.
    assert len(doc.paragraphs) <= 1 # Allow for one default empty paragraph often present in new docs

```
