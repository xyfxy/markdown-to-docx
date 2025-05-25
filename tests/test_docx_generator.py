import pytest
from unittest.mock import MagicMock, patch, call # Using unittest.mock directly
from pathlib import Path
import sys
import os
import logging # For capturing log messages if DocxWriter uses logging for warnings

# Adjust sys.path for local imports
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from docx_generator import (
    DocxWriter, 
    FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, FONT_ENGLISH_NUMBERS,
    FONT_HEADING_H1, FONT_HEADING_H1_FALLBACK,
    FONT_HEADING_H2, FONT_HEADING_H2_FALLBACK,
    FONT_HEADING_H3, FONT_HEADING_H3_FALLBACK,
    DEFAULT_BODY_FONT_SIZE_PT, DEFAULT_LINE_SPACING_PT,
    segment_text_for_font_styling
)
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# --- Fixtures ---

@pytest.fixture
def mock_docx_document_instance(mocker):
    """Provides a fully mocked instance of a docx.Document object."""
    mock_doc_instance = mocker.MagicMock(name="DocumentInstance")
    
    # Mock sections and their properties
    mock_section = mocker.MagicMock(name="Section")
    mock_doc_instance.sections = [mock_section]
    
    # Mock styles and their properties
    mock_normal_style = mocker.MagicMock(name="NormalStyle")
    mock_normal_style.font = mocker.MagicMock(name="NormalStyleFont")
    mock_normal_style.paragraph_format = mocker.MagicMock(name="NormalStyleParagraphFormat")
    mock_doc_instance.styles = {'Normal': mock_normal_style}

    # Mock paragraph and run creation for general use
    # These will be the objects returned when add_paragraph or add_run are called on the document/paragraph
    mock_paragraph = mocker.MagicMock(name="Paragraph")
    mock_paragraph.paragraph_format = mocker.MagicMock(name="ParagraphFormat") # Each paragraph has its own format obj
    mock_doc_instance.add_paragraph.return_value = mock_paragraph
    
    mock_run = mocker.MagicMock(name="Run")
    mock_run.font = mocker.MagicMock(name="RunFont") # Each run has its own font obj
    mock_paragraph.add_run.return_value = mock_run

    return mock_doc_instance

@pytest.fixture
def writer_default(mocker, mock_docx_document_instance):
    """Fixture for a DocxWriter instance with default init params, using a predefined mocked Document instance."""
    # Patch the Document CLASS in the docx_generator module for the duration of this fixture's setup
    # This ensures that when DocxWriter() is called, it uses a MagicMock for self.document
    # that we can then replace with our more detailed mock_docx_document_instance.
    mocker.patch('docx_generator.Document', return_value=mock_docx_document_instance, autospec=True)
    
    writer = DocxWriter()
    # Critical: Ensure the writer instance uses our specific, detailed mock_docx_document_instance
    # This allows assertions on the same object that the fixture prepared.
    writer.document = mock_docx_document_instance 
    return writer


# --- Initialization Tests ---

def test_init_defaults(writer_default): # writer_default already uses mock_docx_document_instance
    assert writer_default.output_docx_path == "output.docx"
    assert writer_default.custom_line_spacing_pt is None
    assert writer_default.custom_font_size_pt is None
    assert writer_default.body_font_size == DEFAULT_BODY_FONT_SIZE_PT

    mock_section = writer_default.document.sections[0]
    assert mock_section.page_width == Cm(21.0)
    assert mock_section.page_height == Cm(29.7)
    assert mock_section.top_margin == Cm(3.7)
    assert mock_section.bottom_margin == Cm(3.5)
    assert mock_section.left_margin == Cm(2.8)
    assert mock_section.right_margin == Cm(2.6)

    mock_normal_style = writer_default.document.styles['Normal']
    assert mock_normal_style.paragraph_format.line_spacing == Pt(DEFAULT_LINE_SPACING_PT)
    assert mock_normal_style.font.size == Pt(DEFAULT_BODY_FONT_SIZE_PT)
    # Check default font names
    assert mock_normal_style.font.name_east_asia == FONT_CHINESE_PRIMARY # or FONT_CHINESE_FALLBACK if primary raised error (not in mock)
    assert mock_normal_style.font.name_ascii == FONT_ENGLISH_NUMBERS


def test_init_custom_options(mocker):
    custom_ls = 30.0
    custom_fs = 18.0
    
    mock_doc_instance = mocker.MagicMock(name="CustomDocInstance")
    mock_doc_instance.sections = [mocker.MagicMock(name="CustomSection")]
    mock_normal_style = mocker.MagicMock(name="CustomNormalStyle", font=mocker.MagicMock(), paragraph_format=mocker.MagicMock())
    mock_doc_instance.styles = {'Normal': mock_normal_style}
    mocker.patch('docx_generator.Document', return_value=mock_doc_instance, autospec=True)

    writer = DocxWriter(output_docx_path="custom.docx", line_spacing_pt=custom_ls, font_size_pt=custom_fs)
    writer.document = mock_doc_instance # Ensure it uses our mock

    assert writer.output_docx_path == "custom.docx"
    assert writer.custom_line_spacing_pt == custom_ls
    assert writer.custom_font_size_pt == custom_fs
    assert writer.body_font_size == custom_fs
    assert mock_normal_style.paragraph_format.line_spacing == Pt(custom_ls)
    assert mock_normal_style.font.size == Pt(custom_fs)


def test_init_with_valid_template(mocker, mock_docx_document_instance):
    mock_template_path = "dummy_template.docx"
    # Patch the Document CLASS, making its constructor return our pre-configured mock instance
    mock_doc_constructor = mocker.patch('docx_generator.Document', return_value=mock_docx_document_instance, autospec=True)
    
    writer = DocxWriter(template_path=mock_template_path)
    
    mock_doc_constructor.assert_called_once_with(mock_template_path)
    # _setup_page and _apply_default_styles are still called, ensure mock handles this.
    assert writer.document == mock_docx_document_instance # Check that the instance from the constructor is used.


def test_init_with_invalid_template_fallback(mocker, caplog): # Use caplog for logging
    invalid_template_path = "non_existent_template.docx"
    
    mock_fallback_doc_instance = mocker.MagicMock(name="FallbackDocInstance")
    mock_fallback_doc_instance.sections = [mocker.MagicMock()]
    mock_fallback_doc_instance.styles = {'Normal': mocker.MagicMock(font=mocker.MagicMock(), paragraph_format=mocker.MagicMock())}

    def side_effect_for_doc_constructor(path=None):
        if path == invalid_template_path:
            # Simulate the actual error that python-docx might raise for a bad template.
            # PackageNotFoundError is one possibility if the path is bad or file corrupted.
            # For simplicity, using a generic Exception.
            raise Exception("Template load error: Not a valid Word file or path.")
        # This is the fallback Document() call
        return mock_fallback_doc_instance

    mock_doc_constructor = mocker.patch('docx_generator.Document', side_effect=side_effect_for_doc_constructor, autospec=True)
    
    # Capture print statements temporarily for this test, as DocxWriter uses print for this warning
    with patch('builtins.print') as mock_print:
        writer = DocxWriter(template_path=invalid_template_path)
        # Check if print was called with a warning message
        warning_found = any("Warning: Could not load template" in str(arg) for arg_list in mock_print.call_args_list for arg in arg_list[0])
        assert warning_found

    assert mock_doc_constructor.call_args_list == [call(invalid_template_path), call()]
    assert writer.document == mock_fallback_doc_instance


# --- Content Method Tests ---

def test_apply_run_font_style_logic(mocker):
    """More focused test for _apply_run_font_style combinations."""
    writer = DocxWriter() # Instance needed to call the method
    mock_run = mocker.MagicMock(font=mocker.MagicMock())

    # Test Chinese font, bold
    writer._apply_run_font_style(mock_run, 'chinese', "TestFontC", "TestFontCFallback", 16.0, bold=True)
    assert mock_run.font.name == "TestFontC"; assert mock_run.font.name_east_asia == "TestFontC"
    assert mock_run.font.name_ascii == FONT_ENGLISH_NUMBERS; assert mock_run.font.size == Pt(16.0)
    assert mock_run.font.bold is True; assert mock_run.font.italic is None # Assuming default for unset is None or False

    # Test English font, italic
    mock_run.reset_mock() # Reset for next call
    mock_run.font.reset_mock()
    writer._apply_run_font_style(mock_run, 'english_number', "TestFontC", "TestFontCFallback", 12.0, italic=True)
    assert mock_run.font.name == FONT_ENGLISH_NUMBERS; assert mock_run.font.name_ascii == FONT_ENGLISH_NUMBERS
    assert mock_run.font.name_east_asia == "TestFontC"; assert mock_run.font.size == Pt(12.0)
    assert mock_run.font.italic is True; assert mock_run.font.bold is None


def test_add_paragraph_formatting_and_font_segmentation(writer_default, mocker):
    test_text = "Text 这是中文 123."
    mock_doc = writer_default.document
    mock_paragraph = mock_doc.add_paragraph.return_value
    
    spy_apply_run_style = mocker.spy(writer_default, '_apply_run_font_style')
    writer_default.add_paragraph(test_text)

    mock_doc.add_paragraph.assert_called_once_with() # No text arg here
    assert mock_paragraph.paragraph_format.first_line_indent == Pt(2 * DEFAULT_BODY_FONT_SIZE_PT)
    
    expected_segments = segment_text_for_font_styling(test_text) # Use the actual segmenter
    assert spy_apply_run_style.call_count == len(expected_segments)
    
    for i, (seg_text, seg_type) in enumerate(expected_segments):
        # Check that add_run was called with this segment's text
        mock_paragraph.add_run.assert_any_call(seg_text)
        # Check that _apply_run_font_style was called with the correct parameters for this segment
        # The run object passed to _apply_run_font_style is the one returned by mock_paragraph.add_run
        # We assume add_run is called sequentially and returns the same mock_run for simplicity in this spy test,
        # or that the spy captures the specific run object.
        # A more robust way is to check properties of the run objects if they were different mocks.
        current_call_args = spy_apply_run_style.call_args_list[i].args
        assert current_call_args[1] == seg_type # font_type
        assert current_call_args[2] == FONT_CHINESE_PRIMARY # base_font_chinese
        assert current_call_args[4] == DEFAULT_BODY_FONT_SIZE_PT # size_pt


def test_add_emphasis_formatting(writer_default, mocker):
    test_text = "Emphasized Text"
    mock_doc = writer_default.document
    mock_paragraph = mock_doc.add_paragraph.return_value
    spy_apply_run_style = mocker.spy(writer_default, '_apply_run_font_style')

    writer_default.add_emphasis(test_text, style='bold')
    
    assert mock_paragraph.paragraph_format.first_line_indent == Pt(2 * DEFAULT_BODY_FONT_SIZE_PT)
    # Assuming single segment for simplicity here
    spy_apply_run_style.assert_called_once()
    call_args = spy_apply_run_style.call_args.args
    assert call_args[4] == DEFAULT_BODY_FONT_SIZE_PT # size_pt
    assert call_args[5] is True # bold=True
    assert call_args[6] is False # italic=False


def test_add_recipient_formatting(writer_default, mocker):
    test_text = "Recipient Name"
    expected_full_text = test_text + "："
    mock_doc = writer_default.document
    mock_paragraph = mock_doc.add_paragraph.return_value
    spy_apply_run_style = mocker.spy(writer_default, '_apply_run_font_style')

    writer_default.add_recipient(test_text)

    assert mock_paragraph.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert mock_paragraph.paragraph_format.first_line_indent == Pt(0)
    
    # Assuming single segment for simplicity in assertion of _apply_run_font_style call
    mock_paragraph.add_run.assert_any_call(expected_full_text) # Check the text with colon
    spy_apply_run_style.assert_called_once() # If expected_full_text is one segment
    call_args = spy_apply_run_style.call_args.args
    assert call_args[4] == DEFAULT_BODY_FONT_SIZE_PT


def test_add_list_item_formatting(writer_default, mocker):
    test_text = "List item text"
    item_prefix = "1. "
    level = 0
    mock_doc = writer_default.document
    mock_paragraph = mock_doc.add_paragraph.return_value
    spy_apply_run_style = mocker.spy(writer_default, '_apply_run_font_style')

    writer_default.add_list_item(test_text, ordered=True, level=level, item_prefix=item_prefix)

    indent_unit_cm = 0.75
    assert mock_paragraph.paragraph_format.left_indent == Cm(indent_unit_cm * (level + 1))
    assert mock_paragraph.paragraph_format.first_line_indent == Cm(-indent_unit_cm)

    # Check prefix run
    mock_paragraph.add_run.assert_any_call(item_prefix)
    # Check text run
    mock_paragraph.add_run.assert_any_call(test_text)
    
    # _apply_run_font_style is called for prefix and for text segments
    assert spy_apply_run_style.call_count >= 1 # At least one for the text, one for prefix if not empty
    
    # Check call for item_prefix
    prefix_call_args = spy_apply_run_style.call_args_list[0].args
    assert prefix_call_args[0] == mock_paragraph.add_run.return_value # The run for prefix
    assert prefix_call_args[1] == 'english_number' # Prefix font_type
    assert prefix_call_args[4] == DEFAULT_BODY_FONT_SIZE_PT

    # Check call for item text (assuming single segment for simplicity)
    text_call_args = spy_apply_run_style.call_args_list[1].args
    assert text_call_args[1] == 'english_number' # Assuming test_text is English for this simple check
    assert text_call_args[4] == DEFAULT_BODY_FONT_SIZE_PT


# --- Heading Formatting Tests ---
# Test H1 (_add_main_heading) was in previous step, slightly adjusted here.
def test_add_main_heading_h1(writer_default, mocker):
    test_title = "主标题"
    mock_doc = writer_default.document
    # Each call to add_paragraph in _add_main_heading will return the same MagicMock instance
    # if not configured otherwise. Let's assume it's the one from the fixture.
    mock_paragraph_title = mock_doc.add_paragraph.return_value 
    spy_apply_run_style = mocker.spy(writer_default, '_apply_run_font_style')

    writer_default.add_heading(test_title, level=1)

    # First add_paragraph call is for the title itself
    assert mock_paragraph_title.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert mock_paragraph_title.paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY
    assert mock_paragraph_title.paragraph_format.line_spacing == Pt(32)
    
    # Check _apply_run_font_style for the title text
    spy_apply_run_style.assert_any_call(mocker.ANY, 'chinese', FONT_HEADING_H1, FONT_HEADING_H1_FALLBACK, 22)
    
    # Check that two empty paragraphs are added after
    # Total add_paragraph calls: 1 for title + 2 for empty lines
    assert mock_doc.add_paragraph.call_count == 3


@pytest.mark.parametrize("level, prefix_text, font_name, fallback_font", [
    (2, "一、 ", FONT_HEADING_H2, FONT_HEADING_H2_FALLBACK),
    (3, "（一） ", FONT_HEADING_H3, FONT_HEADING_H3_FALLBACK),
    (4, "1. ", FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK),
    (5, "（1） ", FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK),
])
def test_add_heading_levels_h2_h5(writer_default, mocker, level, prefix_text, font_name, fallback_font):
    test_text_content = "Heading Content"
    full_expected_text = prefix_text + test_text_content
    
    mock_doc = writer_default.document
    mock_paragraph = mock_doc.add_paragraph.return_value
    spy_apply_run_style = mocker.spy(writer_default, '_apply_run_font_style')

    writer_default.add_heading(test_text_content, level=level)

    mock_doc.add_paragraph.assert_called_once_with() # No text arg
    
    # Verify _apply_run_font_style for segments of `full_expected_text`
    # This assumes prefix is one segment and text_content is another if types differ, or combined if same.
    # For simplicity, let's check if it's called with the correct font and size for the content part.
    # A more robust test would iterate through segment_text_for_font_styling(full_expected_text)
    
    # Example: Check the last call to _apply_run_font_style, assuming it's for 'Heading Content'
    last_call_args = spy_apply_run_style.call_args.args
    assert last_call_args[2] == font_name # base_font_chinese
    assert last_call_args[3] == fallback_font # base_font_chinese_fallback
    assert last_call_args[4] == 16 # size_pt for H2-H5 is 16pt

    # Verify the prefix was part of some run
    # This means add_run was called with text that includes the prefix
    run_texts = [call_args[0][0] for call_args in mock_paragraph.add_run.call_args_list]
    assert any(prefix_text in text for text in run_texts) or prefix_text in "".join(run_texts)


# --- Test Save Method ---
def test_save_method_called(writer_default): # writer_default uses mock_docx_document_instance
    writer_default.save()
    writer_default.document.save.assert_called_once_with(writer_default.output_docx_path)

# --- Test for segment_text_for_font_styling utility --- (already present in previous step)
def test_segment_text_utility(): # Renamed to avoid pytest collection warning if class-based
    assert segment_text_for_font_styling("Hello World") == [("Hello World", "english_number")]
    assert segment_text_for_font_styling("你好世界") == [("你好世界", "chinese")]
    assert segment_text_for_font_styling("Hello 你好 World 世界") == [
        ("Hello ", "english_number"), ("你好", "chinese"),
        (" World ", "english_number"), ("世界", "chinese")
    ]
    assert segment_text_for_font_styling("") == []

```
