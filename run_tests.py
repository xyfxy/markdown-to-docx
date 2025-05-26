import os
from click.testing import CliRunner
from main import markdown_to_docx_cli # Assuming your CLI function is named markdown_to_docx_cli
from docx import Document # For optional checks

# Define file paths (assuming run_tests.py is in the project root)
INPUT_MD = "test_suite.md"
OUTPUT_DOCX = "test_suite_output.docx"
IMAGE_FILE = "sample_image.png" # Used for checking existence

def check_docx_content(docx_path):
    """
    Optional: Performs basic checks on the generated DOCX file.
    """
    if not os.path.exists(docx_path):
        print(f"Error: Output file '{docx_path}' was not generated.")
        return False

    try:
        document = Document(docx_path)
        passed_checks = True
        
        # Check 1: At least one table
        if len(document.tables) > 0:
            print(f"Check PASSED: Document contains {len(document.tables)} table(s).")
        else:
            print("Check FAILED: Document does not contain any tables.")
            passed_checks = False
            
        # Check 2: At least one paragraph (basic sanity check)
        if len(document.paragraphs) > 0:
            print(f"Check PASSED: Document contains {len(document.paragraphs)} paragraph(s).")
        else:
            print("Check FAILED: Document does not contain any paragraphs.")
            passed_checks = False

        # Check 3: Presence of image (simplified check - looks for the image filename in relationships)
        # This is a proxy, as directly finding inline shapes is complex.
        # We assume the image was added if its relationship exists.
        image_found_in_rels = False
        for rel in document.part.rels:
            if document.part.rels[rel].target_ref == f"media/{IMAGE_FILE}": # Default image part name in docx
                 image_found_in_rels = True
                 break
            # Sometimes the image name in rels might be different, e.g. image1.png, image2.png etc.
            # A more robust check might look for any image relationship if only one image is expected.
            # Or, if multiple, check if 'sample_image.png' is part of target_ref if it's renamed by docx
            # For this test, we assume 'sample_image.png' is used as is or is part of the rel target.
            # A simpler check if only one image is expected:
            # if "image" in document.part.rels[rel].target_ref:
            # image_found_in_rels = True; break

        # A more robust check for images if there's only one image expected:
        # Check if there's any image relationship.
        has_any_image_relationship = any(
            "image" in document.part.rels[rel].target_ref.lower() for rel in document.part.rels
        )

        if has_any_image_relationship: # Use this more general check
            print(f"Check PASSED: Document appears to contain an image (relationship found).")
        else:
            print(f"Check FAILED: Document does not appear to contain an image (no image relationship found for '{IMAGE_FILE}' or any image).")
            # To debug, print available image relationships:
            # for rel_id in document.part.rels:
            #     rel = document.part.rels[rel_id]
            #     if "image" in rel.target_ref:
            #         print(f"  Found image rel: {rel.target_ref}")
            passed_checks = False
            
        return passed_checks

    except Exception as e:
        print(f"Error opening or checking DOCX file '{docx_path}': {e}")
        return False

if __name__ == "__main__":
    print(f"Starting test conversion of '{INPUT_MD}' to '{OUTPUT_DOCX}'...")
    
    # Verify that the input markdown file and image exist before running the CLI
    if not os.path.exists(INPUT_MD):
        print(f"Error: Test markdown file '{INPUT_MD}' not found. Please create it first.")
        exit(1)
    if not os.path.exists(IMAGE_FILE):
        print(f"Error: Sample image file '{IMAGE_FILE}' not found. Please create it first.")
        exit(1)

    runner = CliRunner()
    # Added --verbose for more detailed output from the CLI itself during testing
    result = runner.invoke(markdown_to_docx_cli, [INPUT_MD, OUTPUT_DOCX, "--verbose"])

    if result.exit_code == 0:
        print("\nCLI command executed successfully.")
        print(f"Output from CLI:\n{result.output}")
        print(f"Test conversion successful: '{OUTPUT_DOCX}' should be generated.")
        
        print("\n--- Performing basic content checks ---")
        if check_docx_content(OUTPUT_DOCX):
            print("All basic content checks passed.")
            print("Please manually inspect 'test_suite_output.docx' for detailed verification of all elements.")
        else:
            print("Some basic content checks FAILED.")
            print("Please inspect 'test_suite_output.docx' and CLI output for errors.")
    else:
        print("\nCLI command execution FAILED.")
        print(f"Exit Code: {result.exit_code}")
        print(f"Output from CLI:\n{result.output}")
        print("Test conversion failed.")

    # Example of how to run with a template (optional, if you have one)
    # print("\n--- Testing with a template (if you have 'my_template.docx') ---")
    # if os.path.exists("my_template.docx"):
    #     result_template = runner.invoke(markdown_to_docx_cli, [
    #         INPUT_MD, 
    #         "test_suite_output_templated.docx", 
    #         "--template", "my_template.docx",
    #         "--verbose"
    #     ])
    #     if result_template.exit_code == 0:
    #         print("Templated conversion successful: 'test_suite_output_templated.docx' generated.")
    #         check_docx_content("test_suite_output_templated.docx")
    #     else:
    #         print("Templated conversion failed.")
    #         print(result_template.output)
    # else:
    #     print("Skipping template test: 'my_template.docx' not found.")
