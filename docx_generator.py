# docx_generator.py
"""
This module provides the DocxWriter class for creating and populating .docx files
based on structured content, with specific formatting for Chinese documents.
"""
import re
import typing # Added for Optional type hint

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
# from docx.enum.style import WD_STYLE_TYPE # For creating custom styles if needed later

# --- Font Name Constants ---
FONT_CHINESE_PRIMARY = "仿宋_GB2312"  # Windows系统中的仿宋字体名称
FONT_CHINESE_FALLBACK = "FangSong"    # 备用仿宋字体名称
FONT_CHINESE_FALLBACK2 = "SimSun"     # 最终回退字体
FONT_ENGLISH_NUMBERS = "Times New Roman"
FONT_HEADING_H1 = "仿宋_GB2312"  # 改为仿宋字体
FONT_HEADING_H1_FALLBACK = "FangSong"  # 仿宋回退字体
FONT_HEADING_H1_FALLBACK2 = "SimSun"  # 最终回退
FONT_HEADING_H2 = "黑体"
FONT_HEADING_H2_FALLBACK = "SimHei"
FONT_HEADING_H3 = "楷体_GB2312"       # Windows系统中的楷体字体名称
FONT_HEADING_H3_FALLBACK = "KaiTi"

# Default PRD values - 3号字体约等于16pt
DEFAULT_BODY_FONT_SIZE_PT = 16.0
DEFAULT_LINE_SPACING_PT = 27.0

# 字号对照表（中文字号到磅值的转换）
FONT_SIZE_MAPPING = {
    "二号": 22.0,  # H1标题
    "三号": 16.0,  # 正文和其他标题
}

# --- Additional Font Name Constants for Code ---
FONT_CODE_ENGLISH = "Courier New"
FONT_CODE_CHINESE = "SimSun" # Or NSimSun, SimSun-ExtB
# FONT_CODE_CHINESE = "NSimSun" # Alternative

# 中文数字转换
CHINESE_NUMBERS = ["", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]

def convert_to_chinese_number(num):
    """将阿拉伯数字转换为中文数字"""
    if num <= 0 or num > 10:
        return str(num)  # 超出范围返回原数字
    return CHINESE_NUMBERS[num]

def is_chinese_char(char):
    return '\u4e00' <= char <= '\u9fff'

def segment_text_for_font_styling(text: str) -> list[tuple[str, str]]:
    segments = []
    if not text:
        return segments
    current_segment = ""
    current_type = 'chinese' if is_chinese_char(text[0]) else 'english_number'
    for char in text:
        char_type = 'chinese' if is_chinese_char(char) else 'english_number'
        if char_type == current_type:
            current_segment += char
        else:
            segments.append((current_segment, current_type))
            current_segment = char
            current_type = char_type
    if current_segment:
        segments.append((current_segment, current_type))
    return segments


class DocxWriter:
    def __init__(self, 
                 output_docx_path: str = "output.docx",
                 line_spacing_pt: typing.Optional[float] = None,
                 font_size_pt: typing.Optional[float] = None,
                 template_path: typing.Optional[str] = None): # Added template_path
        
        self.output_docx_path = output_docx_path
        # If a template is provided, python-docx uses it as the basis for the new document.
        # Styles and some properties can be inherited from the template.
        if template_path:
            try:
                self.document = Document(template_path)
                print(f"Info: Using document template from '{template_path}'")
            except Exception as e:
                print(f"Warning: Could not load template '{template_path}'. Using default. Error: {e}")
                self.document = Document()
        else:
            self.document = Document()

        # Store custom formatting options
        self.custom_line_spacing_pt = line_spacing_pt
        self.custom_font_size_pt = font_size_pt
        
        # Body font size to use (either custom or default)
        self.body_font_size = self.custom_font_size_pt if self.custom_font_size_pt is not None else DEFAULT_BODY_FONT_SIZE_PT

        # 层次编号计数器
        self.heading_counters = {
            2: 0,  # 二级标题：一、二、三...
            3: 0,  # 三级标题：（一）（二）（三）...
            4: 0,  # 四级标题：1. 2. 3. ...
            5: 0   # 五级标题：（1）（2）（3）...
        }

        self._setup_page() # Page setup should apply even with a template, might override template settings
        self._apply_default_styles()


    def _setup_page(self):
        section = self.document.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

    def _apply_default_styles(self):
        style = self.document.styles['Normal']
        p_format = style.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Apply custom line spacing if provided, else default
        actual_line_spacing = self.custom_line_spacing_pt if self.custom_line_spacing_pt is not None else DEFAULT_LINE_SPACING_PT
        p_format.line_spacing = Pt(actual_line_spacing)
        
        font_style = style.font
        # 设置正文字体为仿宋
        self._set_font_with_fallback(font_style, FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, FONT_CHINESE_FALLBACK2)
        
        # Apply custom body font size to the 'Normal' style's font if provided.
        # This helps ensure consistency if paragraphs are added without explicit run styling.
        font_style.size = Pt(self.body_font_size)

    def _set_font_with_fallback(self, font_obj, primary_font: str, fallback_font: str, fallback_font2: str = None):
        """设置字体，包含多级回退机制，使用XML方式确保WPS兼容性"""
        fonts_to_try = [primary_font, fallback_font]
        if fallback_font2:
            fonts_to_try.append(fallback_font2)
        fonts_to_try.append("SimSun")  # 最终回退
        
        selected_font = None
        for font_name in fonts_to_try:
            try:
                # 先尝试设置常规属性
                font_obj.name = font_name
                selected_font = font_name
                break
            except Exception:
                continue
        
        # 使用XML方式设置东亚字体，确保WPS识别
        if selected_font and hasattr(font_obj, 'element') and hasattr(font_obj.element, 'rPr'):
            try:
                from docx.oxml import OxmlElement
                rPr = font_obj.element.rPr
                
                # 获取或创建rFonts元素
                rFonts = rPr.rFonts
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                
                # 设置字体
                rFonts.set(qn('w:eastAsia'), selected_font)
                rFonts.set(qn('w:ascii'), FONT_ENGLISH_NUMBERS)
                rFonts.set(qn('w:hAnsi'), FONT_ENGLISH_NUMBERS)
                
            except Exception as e:
                # 回退到标准方式
                try:
                    font_obj.name_east_asia = selected_font
                    font_obj.name_ascii = FONT_ENGLISH_NUMBERS
                except Exception:
                    pass

    def _apply_run_font_style(self, run, font_type: str, base_font_chinese: str, base_font_chinese_fallback: str, size_pt: float, bold: bool = False, italic: bool = False):
        run.font.size = Pt(size_pt)
        if font_type == 'chinese':
            # 为中文字体提供更好的回退机制
            if base_font_chinese == FONT_CHINESE_PRIMARY:
                self._set_font_with_fallback(run.font, base_font_chinese, base_font_chinese_fallback, FONT_CHINESE_FALLBACK2)
            elif base_font_chinese == FONT_HEADING_H1:
                self._set_font_with_fallback(run.font, base_font_chinese, FONT_HEADING_H1_FALLBACK, FONT_HEADING_H1_FALLBACK2)
            # Check if it's a code font
            elif base_font_chinese == FONT_CODE_CHINESE:
                 self._set_font_with_fallback(run.font, FONT_CODE_CHINESE, "SimSun") # Fallback for Chinese code font
            else:
                self._set_font_with_fallback(run.font, base_font_chinese, base_font_chinese_fallback)
        else: # 'english_number' or potentially english part of code
            # If the base_font_chinese is actually the code font, then English part should also be code font
            if base_font_chinese == FONT_CODE_CHINESE: # Indicates we are in a code segment
                run.font.name = FONT_CODE_ENGLISH
            else: # Standard English/number font
                run.font.name = FONT_ENGLISH_NUMBERS
            
            # 使用XML方式设置字体，确保WPS兼容性
            if hasattr(run.font, 'element') and hasattr(run.font.element, 'rPr'):
                try:
                    from docx.oxml import OxmlElement
                    rPr = run.font.element.rPr
                    
                    # 获取或创建rFonts元素
                    rFonts = rPr.rFonts
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.append(rFonts)
                    
                    current_ascii_font = FONT_CODE_ENGLISH if base_font_chinese == FONT_CODE_CHINESE else FONT_ENGLISH_NUMBERS
                    current_east_asia_font = FONT_CODE_CHINESE if base_font_chinese == FONT_CODE_CHINESE else base_font_chinese

                    rFonts.set(qn('w:ascii'), current_ascii_font)
                    rFonts.set(qn('w:hAnsi'), current_ascii_font)
                    # 东亚字体使用中文字体，确保混排一致性
                    if current_east_asia_font == FONT_CHINESE_PRIMARY:
                        rFonts.set(qn('w:eastAsia'), FONT_CHINESE_PRIMARY)
                    else:
                        rFonts.set(qn('w:eastAsia'), current_east_asia_font)
                        
                except Exception:
                    # 回退到标准方式
                    try:
                        current_ascii_font = FONT_CODE_ENGLISH if base_font_chinese == FONT_CODE_CHINESE else FONT_ENGLISH_NUMBERS
                        current_east_asia_font = FONT_CODE_CHINESE if base_font_chinese == FONT_CODE_CHINESE else base_font_chinese
                        run.font.name_ascii = current_ascii_font
                        run.font.name_east_asia = current_east_asia_font
                    except Exception:
                        run.font.name_east_asia = base_font_chinese_fallback # Fallback for eastAsia if specific code font fails
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True

    def add_recipient(self, text: str):
        full_recipient_text = text + "："
        p = self.document.add_paragraph()
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_format.first_line_indent = Pt(0)
        segments = segment_text_for_font_styling(full_recipient_text)
        for segment_text, font_type in segments:
            run = p.add_run(segment_text)
            self._apply_run_font_style(run, font_type, FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, self.body_font_size)
        return p

    def add_paragraph(self, text: str):
        p = self.document.add_paragraph()
        p_format = p.paragraph_format
        p_format.first_line_indent = Pt(32) # Standard 2-char indent for 16pt font
        # Adjust indent if font size changes significantly?
        # For now, 32pt is based on 16pt. If body_font_size is e.g. 12pt, indent might be 24pt.
        # A common way is 2 * body_font_size.
        p_format.first_line_indent = Pt(2 * self.body_font_size)


        segments = segment_text_for_font_styling(text)
        for segment_text, font_type in segments:
            run = p.add_run(segment_text)
            self._apply_run_font_style(run, font_type, FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, self.body_font_size)
        return p

    def add_emphasis(self, text: str, style: str):
        p = self.document.add_paragraph()
        p_format = p.paragraph_format
        p_format.first_line_indent = Pt(2 * self.body_font_size) # Consistent indent
        is_bold = (style == 'bold')
        is_italic = (style == 'italic')
        segments = segment_text_for_font_styling(text)
        for segment_text, font_type in segments:
            run = p.add_run(segment_text)
            self._apply_run_font_style(run, font_type, FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, self.body_font_size, bold=is_bold, italic=is_italic)
        return p

    # --- Heading Methods ---
    # Headings retain their PRD-specified fixed sizes, not affected by CLI --font-size for body text.
    def _add_main_heading(self, text: str):
        """添加主标题（H1）- 二号方正小标宋简体，居中，不加粗，行距32磅"""
        p = self.document.add_paragraph()
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p_format.line_spacing = Pt(32) # Fixed line spacing for H1
        segments = segment_text_for_font_styling(text)
        for segment_text, font_type in segments:
            run = p.add_run(segment_text)
            # 使用二号字体（22pt），不加粗
            self._apply_run_font_style(run, font_type, FONT_HEADING_H1, FONT_HEADING_H1_FALLBACK, FONT_SIZE_MAPPING["二号"], bold=False)
        # 标题下空两行
        self.document.add_paragraph()
        self.document.add_paragraph()

    def _add_h2(self, text_content: str):
        """添加二级标题 - 黑体，使用"一、" """
        self.heading_counters[2] += 1
        # 重置下级标题计数器
        self.heading_counters[3] = 0
        self.heading_counters[4] = 0
        self.heading_counters[5] = 0
        
        chinese_num = convert_to_chinese_number(self.heading_counters[2])
        full_text = f"{chinese_num}、{text_content}"
        p = self.document.add_paragraph()
        segments = segment_text_for_font_styling(full_text)
        for segment_text, font_type in segments:
            run = p.add_run(segment_text)
            # 使用三号黑体
            self._apply_run_font_style(run, font_type, FONT_HEADING_H2, FONT_HEADING_H2_FALLBACK, FONT_SIZE_MAPPING["三号"])

    def _add_h3(self, text_content: str):
        """添加三级标题 - 楷体，使用"（一）" """
        self.heading_counters[3] += 1
        # 重置下级标题计数器
        self.heading_counters[4] = 0
        self.heading_counters[5] = 0
        
        chinese_num = convert_to_chinese_number(self.heading_counters[3])
        full_text = f"（{chinese_num}）{text_content}"
        p = self.document.add_paragraph()
        segments = segment_text_for_font_styling(full_text)
        for segment_text, font_type in segments:
            run = p.add_run(segment_text)
            # 使用三号楷体
            self._apply_run_font_style(run, font_type, FONT_HEADING_H3, FONT_HEADING_H3_FALLBACK, FONT_SIZE_MAPPING["三号"])

    def _add_h4(self, text_content: str):
        """添加四级标题 - 仿宋体，使用"1." """
        self.heading_counters[4] += 1
        # 重置下级标题计数器
        self.heading_counters[5] = 0
        
        full_text = f"{self.heading_counters[4]}.{text_content}"
        p = self.document.add_paragraph()
        segments = segment_text_for_font_styling(full_text)
        for segment_text, font_type in segments:
            run = p.add_run(segment_text)
            # 使用三号仿宋体
            self._apply_run_font_style(run, font_type, FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, FONT_SIZE_MAPPING["三号"])
            
    def _add_h5(self, text_content: str):
        """添加五级标题 - 仿宋体，使用"（1）" """
        self.heading_counters[5] += 1
        
        full_text = f"（{self.heading_counters[5]}）{text_content}"
        p = self.document.add_paragraph()
        segments = segment_text_for_font_styling(full_text)
        for segment_text, font_type in segments:
            run = p.add_run(segment_text)
            # 使用三号仿宋体
            self._apply_run_font_style(run, font_type, FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, FONT_SIZE_MAPPING["三号"])

    def add_heading(self, text: str, level: int):
        if level == 1: self._add_main_heading(text)
        elif level == 2: self._add_h2(text)
        elif level == 3: self._add_h3(text)
        elif level == 4: self._add_h4(text)
        elif level == 5: self._add_h5(text)
        else: self.add_paragraph(text) # Fallback with body font size

    def add_list_item(self, text: str, ordered: bool, level: int = 0, item_prefix: str = ""):
        p = self.document.add_paragraph()
        p_format = p.paragraph_format
        indent_unit_cm = 0.75 
        p_format.left_indent = Cm(indent_unit_cm * (level + 1))
        p_format.first_line_indent = Cm(-indent_unit_cm)
        
        if item_prefix:
            prefix_run = p.add_run(item_prefix)
            self._apply_run_font_style(prefix_run, 'english_number', FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, self.body_font_size)

        segments = segment_text_for_font_styling(text)
        for segment_text, font_type in segments:
            run = p.add_run(segment_text)
            self._apply_run_font_style(run, font_type, FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, self.body_font_size)
        return p

    # --- New element handling methods ---

    def add_table(self, table_data: dict):
        """
        Adds a table to the document based on table_data.
        table_data format: {'headers': [['H1', 'H2']], 'rows': [['R1C1', 'R1C2'], ['R2C1', 'R2C2']]}
        Headers can be multi-row.
        """
        num_header_rows = len(table_data.get('headers', []))
        first_data_row_idx = num_header_rows
        
        # Calculate total rows and columns
        # Columns: max number of cells in any header row or data row
        all_rows_content = table_data.get('headers', []) + table_data.get('rows', [])
        if not all_rows_content:
            return # Don't add an empty table
            
        num_cols = 0
        for row_content in all_rows_content:
            if isinstance(row_content, list): # row_content is a list of cells (strings or dicts)
                num_cols = max(num_cols, len(row_content))
            # If row_content is not a list, it's an error in table_data structure,
            # but max() will fail. For now, assume correct structure from md_parser.
            
        if num_cols == 0 and not all_rows_content: # No data at all
             return
        if num_cols == 0 and all_rows_content: # Has rows, but no cells (e.g. [[''], ['']])
             num_cols = 1 # Default to 1 column if rows exist but are empty lists of cells


        total_rows = len(all_rows_content)
        if total_rows == 0:
            return

        try:
            table = self.document.add_table(rows=total_rows, cols=num_cols)
            table.style = 'Table Grid' # Apply a basic style with borders
        except Exception as e:
            print(f"Error creating table structure: {e}")
            # Add a paragraph indicating failure to create table
            self.add_paragraph(f"[Error creating table: {e}]")
            return

        # Populate header rows
        for r_idx, header_row_content in enumerate(table_data.get('headers', [])):
            if r_idx >= total_rows: continue # Should not happen with correct total_rows calc
            row_cells = table.rows[r_idx].cells
            for c_idx, cell_text_or_dict in enumerate(header_row_content):
                if c_idx >= num_cols: continue # Should not happen
                cell_text = cell_text_or_dict.get('text') if isinstance(cell_text_or_dict, dict) else cell_text_or_dict
                
                # Clear default paragraph in cell
                cell_paragraph = row_cells[c_idx].paragraphs[0]
                while len(cell_paragraph.runs):
                    cell_paragraph.runs[0]._r.getparent().remove(cell_paragraph.runs[0]._r)

                segments = segment_text_for_font_styling(str(cell_text)) # Ensure text is string
                for segment_text, font_type in segments:
                    run = row_cells[c_idx].add_paragraph().add_run(segment_text) # Add new paragraph for run
                    # For headers, use body font size but make it bold
                    self._apply_run_font_style(run, font_type, FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, self.body_font_size, bold=True)
                # Remove the initial empty paragraph if content was added
                if cell_text and len(row_cells[c_idx].paragraphs) > 1:
                     p_to_remove = row_cells[c_idx].paragraphs[0]._element
                     p_to_remove.getparent().remove(p_to_remove)


        # Populate data rows
        for r_idx, data_row_content in enumerate(table_data.get('rows', [])):
            actual_r_idx = first_data_row_idx + r_idx
            if actual_r_idx >= total_rows: continue
            row_cells = table.rows[actual_r_idx].cells
            for c_idx, cell_text_or_dict in enumerate(data_row_content):
                if c_idx >= num_cols: continue
                cell_text = cell_text_or_dict.get('text') if isinstance(cell_text_or_dict, dict) else cell_text_or_dict

                # Clear default paragraph in cell
                cell_paragraph = row_cells[c_idx].paragraphs[0]
                while len(cell_paragraph.runs):
                    cell_paragraph.runs[0]._r.getparent().remove(cell_paragraph.runs[0]._r)

                segments = segment_text_for_font_styling(str(cell_text))
                for segment_text, font_type in segments:
                    run = row_cells[c_idx].add_paragraph().add_run(segment_text)
                    self._apply_run_font_style(run, font_type, FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, self.body_font_size)
                # Remove the initial empty paragraph if content was added
                if cell_text and len(row_cells[c_idx].paragraphs) > 1:
                     p_to_remove = row_cells[c_idx].paragraphs[0]._element
                     p_to_remove.getparent().remove(p_to_remove)
        
        # Add a blank paragraph after the table for spacing, if desired
        # self.document.add_paragraph()


    def add_code_block(self, code_text: str, language: typing.Optional[str] = None):
        """Adds a code block with monospaced font and optional background color."""
        p = self.document.add_paragraph()
        # Optional: Set background color for the paragraph (requires OXML)
        # try:
        #     from docx.oxml import OxmlElement
        #     from docx.oxml.ns import qn
        #     shd = OxmlElement('w:shd')
        #     shd.set(qn('w:val'), 'clear')
        #     shd.set(qn('w:color'), 'auto')
        #     shd.set(qn('w:fill'), 'F0F0F0') # Light gray
        #     p._p.get_or_add_pPr().append(shd)
        # except ImportError:
        #     print("Warning: Could not import OxmlElement for code block background.")

        # Code blocks should not have first line indent.
        p.paragraph_format.first_line_indent = Pt(0)

        segments = segment_text_for_font_styling(code_text)
        for segment_text, font_type in segments:
            run = p.add_run(segment_text)
            # Use FONT_CODE_CHINESE as the 'base_font_chinese' argument
            # _apply_run_font_style will then use FONT_CODE_ENGLISH for 'english_number' type segments
            self._apply_run_font_style(run, font_type, FONT_CODE_CHINESE, FONT_CODE_CHINESE, self.body_font_size)
            # Note: FONT_CODE_CHINESE is used as fallback for itself, assuming it's a reliable font.

    def add_inline_code(self, text: str, paragraph_obj):
        """Adds an inline code run to an existing paragraph object."""
        if not paragraph_obj:
            # Fallback to creating a new paragraph if none is provided (though not ideal for inline)
            paragraph_obj = self.document.add_paragraph()
            paragraph_obj.paragraph_format.first_line_indent = Pt(0) # Assume no indent if new para

        segments = segment_text_for_font_styling(text)
        for segment_text, font_type in segments:
            run = paragraph_obj.add_run(segment_text)
            self._apply_run_font_style(run, font_type, FONT_CODE_CHINESE, FONT_CODE_CHINESE, self.body_font_size)


    def add_blockquote(self, text: str):
        """Adds a blockquote with left indentation and italic style."""
        p = self.document.add_paragraph()
        p_format = p.paragraph_format
        p_format.left_indent = Cm(1.25)
        # Blockquotes typically don't have first-line indent beyond the block indent
        p_format.first_line_indent = Pt(0) 

        segments = segment_text_for_font_styling(text)
        for segment_text, font_type in segments:
            run = p.add_run(segment_text)
            self._apply_run_font_style(run, font_type, FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, self.body_font_size, italic=True)


    def add_hyperlink_paragraph(self, text: str, url: str):
        """Adds a paragraph containing a single hyperlink."""
        # This is a simplified version. For complex scenarios (hyperlink within existing text),
        # direct OXML manipulation is needed. python-docx > 0.8.x has add_hyperlink on runs.
        # Assuming a version that supports paragraph.add_hyperlink (which might not exist directly)
        # Fallback: create a paragraph and style the text to look like a link.
        # Actual hyperlink functionality requires more complex OXML or a newer python-docx.
        
        # The python-docx library does not have a direct paragraph.add_hyperlink method.
        # It's add_run().add_hyperlink() or more complex for full paragraph links.
        # For now, we create a styled paragraph.
        # A common way to add actual hyperlinks is by using a character style named 'Hyperlink'
        # and then applying it, or direct OXML.
        
        # Simplest approach: create a run, style it blue and underlined.
        # This won't be a clickable hyperlink in Word without more work.
        # However, the PRD implies creating a new paragraph for the link.

        p = self.document.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0) # Links usually don't have indent

        # For actual hyperlink, one would typically do:
        # from docx.oxml.shared import ns
        # hyperlink = OxmlElement('w:hyperlink')
        # hyperlink.set(ns.qn('r:id'), self.document.part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True))
        # run_element = OxmlElement('w:r')
        # run_element.append(OxmlElement('w:rPr'))
        # run_element.rPr.append(OxmlElement('w:rStyle')).set(ns.qn('w:val'), 'Hyperlink') # Assumes 'Hyperlink' style exists
        # run_element.append(OxmlElement('w:t')).text = text
        # hyperlink.append(run_element)
        # p._p.append(hyperlink)
        # This is complex. A simpler visual cue:

        segments = segment_text_for_font_styling(text)
        for segment_text, font_type in segments:
            run = p.add_run(segment_text)
            self._apply_run_font_style(run, font_type, FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, self.body_font_size)
            run.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1) # Standard blue link color
            run.font.underline = True
        # Add the URL in parentheses after the link text for clarity, since it's not a real hyperlink.
        if url:
            url_run = p.add_run(f" ({url})")
            self._apply_run_font_style(url_run, 'english_number', FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, self.body_font_size)
            url_run.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
            url_run.font.underline = True


    def add_image(self, src: str, alt: typing.Optional[str] = None, width_cm: typing.Optional[float] = None):
        """Adds an image from a local path."""
        import os # For path validation
        if src.startswith(('http://', 'https://')):
            print(f"Warning: Skipping image from URL (not supported): {src}")
            self.add_paragraph(f"[Image skipped: {src} (URL not supported). Alt text: {alt or 'N/A'}]")
            return

        if not os.path.exists(src) or not os.path.isfile(src):
            print(f"Warning: Image file not found or is not a file: {src}")
            self.add_paragraph(f"[Image skipped: File not found at {src}. Alt text: {alt or 'N/A'}]")
            return
        
        try:
            if width_cm:
                self.document.add_picture(src, width=Cm(width_cm))
            else:
                self.document.add_picture(src) # Add with default size
            # Optionally add alt text as a caption below the image
            if alt:
                # Create a new paragraph for the caption, centered, smaller font.
                p_caption = self.document.add_paragraph()
                p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Caption specific font size (e.g., smaller than body)
                caption_font_size = self.body_font_size * 0.8 
                if caption_font_size < 9: caption_font_size = 9 # Minimum size

                segments = segment_text_for_font_styling(alt)
                for segment_text, font_type in segments:
                    run = p_caption.add_run(segment_text)
                    # Using primary body font for caption, but smaller and italic
                    self._apply_run_font_style(run, font_type, FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, caption_font_size, italic=True)
        except Exception as e:
            print(f"Error adding image {src}: {e}")
            self.add_paragraph(f"[Error adding image: {src}. Alt text: {alt or 'N/A'}. Error: {e}]")


    def add_horizontal_rule(self):
        """Adds a horizontal rule using a centered paragraph with '---'."""
        p = self.document.add_paragraph("---")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Ensure it doesn't have first line indent
        p.paragraph_format.first_line_indent = Pt(0)
        # Optional: Make the font size smaller for the rule or change color to gray
        # For now, it will use the default paragraph font settings.


    def save(self):
        try:
            self.document.save(self.output_docx_path)
        except Exception as e:
            print(f"Error saving document to {self.output_docx_path}: {e}")
            raise

# Example usage (for testing, can be removed or commented out)
if __name__ == '__main__':
    writer = DocxWriter("test_document_new_elements.docx")
    writer.add_heading("主标题示例", 1)
    writer.add_recipient("收件人测试")
    writer.add_paragraph("这是一个常规段落，包含中文和 English text, with numbers 12345.")
    
    # Test emphasis
    writer.add_emphasis("这是粗体强调的文本。", "bold")
    writer.add_emphasis("这是斜体强调的文本。", "italic")

    # Test headings
    writer.add_heading("二级标题一", 2)
    writer.add_paragraph("二级标题下的段落。")
    writer.add_heading("三级标题A", 3)
    writer.add_paragraph("三级标题下的段落。")
    writer.add_heading("四级标题i", 4)
    writer.add_paragraph("四级标题下的段落。")
    writer.add_heading("五级标题a", 5)
    writer.add_paragraph("五级标题下的段落。")

    # Test Lists (existing functionality)
    writer.add_list_item("无序列表项1，层级0。", ordered=False, level=0, item_prefix="• ")
    writer.add_list_item("无序列表项2，层级0。", ordered=False, level=0, item_prefix="• ")
    writer.add_list_item("有序列表项A，层级0。", ordered=True, level=0, item_prefix="A. ")
    writer.add_list_item("有序列表项B，层级0。", ordered=True, level=0, item_prefix="B. ")

    # Test Table
    table_content = {
        'headers': [['姓名 (Name)', '年龄 (Age)', '城市 (City)']],
        'rows': [
            ['张三 (Zhang San)', '30', '北京 (Beijing)'],
            ['李四 (Li Si)', '25', '上海 (Shanghai)'],
            [{'text': '王五 (Wang Wu) Code: `Test`', 'is_header': False}, '35', '深圳 (Shenzhen)']
        ]
    }
    writer.add_table(table_content)
    writer.add_paragraph("表格后的段落。")

    # Test Code Block
    writer.add_code_block("def greet(name):\n  print(f\"Hello, {name}!\")\n\n# 这是中文注释\ngreet(\"世界\")", language="python")
    writer.add_paragraph("代码块后的段落。")

    # Test Inline Code (using the new add_inline_code method with a paragraph object)
    p_for_inline = writer.add_paragraph("这是一个段落，它包含") # Get the paragraph object
    writer.add_inline_code(" inline_code_example() ", p_for_inline)
    p_for_inline.add_run(" 和一些常规文本。") # Add more text to the same paragraph
    writer._apply_run_font_style(p_for_inline.runs[-1], 'chinese', FONT_CHINESE_PRIMARY, FONT_CHINESE_FALLBACK, writer.body_font_size)


    # Test Blockquote
    writer.add_blockquote("这是一个块引用示例，它应该被缩进并且是斜体的。This is a blockquote example, it should be indented and italicized.")
    
    # Test Hyperlink Paragraph
    writer.add_hyperlink_paragraph("访问我们的网站 Google", "https://www.google.com")

    # Test Horizontal Rule
    writer.add_horizontal_rule()
    
    # Test Image (Create a dummy image file for testing)
    # You'll need to create a dummy image file named 'dummy_image.png' in the same directory, or provide a valid path.
    # try:
    #     from PIL import Image
    #     img = Image.new('RGB', (100, 50), color = 'red')
    #     img.save("dummy_image.png")
    #     writer.add_image("dummy_image.png", alt="一个红色矩形图片 (A red rectangle image)", width_cm=5)
    # except ImportError:
    #     print("PIL/Pillow not installed, skipping dummy image creation for test.")
    #     writer.add_paragraph("[测试图片被跳过，因为Pillow未安装]")
    # except Exception as e:
    #     print(f"Error creating dummy image: {e}")
    #     writer.add_paragraph(f"[创建测试图片失败: {e}]")
        
    writer.add_paragraph("文档末尾。End of document.")
    writer.save()
    print(f"Test document 'test_document_new_elements.docx' generated successfully.")
